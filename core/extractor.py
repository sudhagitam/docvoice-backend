"""
DocVoice AI – Document Extractor
Handles PDF and DOCX text extraction with automatic OCR fallback.
Supports Telugu, English, and 10 other languages.
"""

import io
import logging
import re
from pathlib import Path

logger = logging.getLogger("docvoice.extractor")

# Tesseract language codes for OCR
TESSERACT_LANG_MAP = {
    "en": "eng",
    "te": "tel",  # Telugu
    "hi": "hin",
    "es": "spa",
    "fr": "fra",
    "de": "deu",
    "it": "ita",
    "pt": "por",
    "ar": "ara",
    "ru": "rus",
    "zh": "chi_sim",
    "ja": "jpn",
    "ko": "kor",
}

# Default OCR language chain — eng+tel gives best results for mixed docs
DEFAULT_OCR_LANG = "eng+tel"


class DocumentExtractor:
    """Extracts plain text from PDF and DOCX files."""

    SUPPORTED = {".pdf", ".docx"}

    def extract(self, file_bytes: bytes, filename: str, lang: str = "en") -> str:
        """
        Extract text from file_bytes.
        Automatically uses OCR for scanned PDFs.

        Args:
            file_bytes: Raw file content
            filename:   Original filename
            lang:       Language hint for OCR (BCP-47 code e.g. 'en', 'te')
        """
        from core.exceptions import (
            CorruptedFileError,
            EmptyDocumentError,
            UnsupportedFileTypeError,
        )

        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED:
            raise UnsupportedFileTypeError(ext)

        logger.info("Extracting '%s' (%d bytes) lang=%s", filename, len(file_bytes), lang)

        try:
            if ext == ".pdf":
                text = self._extract_pdf(file_bytes, lang)
            else:
                text = self._extract_docx(file_bytes)
        except (UnsupportedFileTypeError, EmptyDocumentError, CorruptedFileError):
            raise
        except Exception as exc:
            logger.exception("Unexpected extraction error")
            raise CorruptedFileError(str(exc)) from exc

        text = self._clean(text)

        if not text.strip():
            raise EmptyDocumentError()

        logger.info("Extracted %d characters from '%s'", len(text), filename)
        return text

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _extract_pdf(self, data: bytes, lang: str = "en") -> str:
        from core.exceptions import CorruptedFileError

        try:
            import PyPDF2
        except ImportError as exc:
            raise CorruptedFileError("PyPDF2 not installed") from exc

        try:
            reader = PyPDF2.PdfReader(io.BytesIO(data))
        except Exception as exc:
            raise CorruptedFileError(f"Cannot parse PDF: {exc}") from exc

        if reader.is_encrypted:
            raise CorruptedFileError("Password-protected PDFs are not supported.")

        parts: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                logger.warning("Failed to extract page %d; skipping.", i + 1)

        text = "\n".join(parts)
        total_pages = len(reader.pages)
        avg_chars = len(text.strip()) / max(total_pages, 1)

        # Auto-detect scanned PDF: less than 50 chars per page average
        if avg_chars < 50:
            logger.info(
                "Sparse text detected (%.0f chars/page) — running OCR (lang=%s)",
                avg_chars, lang,
            )
            ocr_text = self._ocr_pdf(data, lang)
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text

        return text

    def _ocr_pdf(self, data: bytes, lang: str = "en") -> str:
        """OCR using pdf2image + pytesseract with language support."""
        try:
            import pytesseract
            from pdf2image import convert_from_bytes

            # Build tesseract language string
            tess_lang = self._get_tess_lang(lang)
            logger.info("Starting OCR: lang=%s tesseract=%s", lang, tess_lang)

            images = convert_from_bytes(
                data,
                dpi=200,
                fmt="jpeg",
                thread_count=2,
            )

            logger.info("OCR: %d pages to process", len(images))

            parts = []
            for i, img in enumerate(images):
                logger.debug("OCR page %d/%d", i + 1, len(images))
                page_text = pytesseract.image_to_string(img, lang=tess_lang)
                parts.append(page_text)

            result = "\n".join(parts)
            logger.info("OCR complete: %d characters", len(result))
            return result

        except ImportError as exc:
            logger.warning("OCR libraries not available: %s", exc)
            return ""
        except Exception as exc:
            logger.warning("OCR failed: %s", exc)
            return ""

    def _get_tess_lang(self, lang: str) -> str:
        """
        Convert BCP-47 language code to tesseract lang string.
        Always includes English as fallback for mixed documents.
        Telugu always paired with English for best results.
        """
        tess = TESSERACT_LANG_MAP.get(lang, "eng")

        # Telugu: always pair with English (mixed docs common)
        if tess == "tel":
            return "tel+eng"

        # Hindi: pair with English
        if tess == "hin":
            return "hin+eng"

        # For other non-English languages, add eng as fallback
        if tess != "eng":
            return f"{tess}+eng"

        return tess

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def _extract_docx(self, data: bytes) -> str:
        from core.exceptions import CorruptedFileError

        try:
            from docx import Document
        except ImportError as exc:
            raise CorruptedFileError("python-docx not installed") from exc

        try:
            doc = Document(io.BytesIO(data))
        except Exception as exc:
            raise CorruptedFileError(f"Cannot parse DOCX: {exc}") from exc

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        return "\n".join(paragraphs)

    # ── Helpers ───────────────────────────────────────────────────────────────

    @staticmethod
    def _clean(text: str) -> str:
        text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]{2,}", " ", text)
        return text.strip()
